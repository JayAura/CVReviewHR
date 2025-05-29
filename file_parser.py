import os
import logging
import re
import io
from typing import Tuple, Optional

# Configure logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

def parse_cv_file(file_path: str) -> Tuple[bool, str]:
    """
    Parse the content of a CV file based on its extension
    
    Args:
        file_path: Path to the CV file
        
    Returns:
        Tuple of (success, content or error message)
    """
    try:
        file_ext = os.path.splitext(file_path)[1].lower()
        
        # PDF parsing
        if file_ext == '.pdf':
            return parse_pdf(file_path)
        
        # DOCX parsing
        elif file_ext == '.docx':
            return parse_docx(file_path)
        
        # DOC parsing
        elif file_ext == '.doc':
            return parse_doc(file_path)
        
        # TXT parsing
        elif file_ext == '.txt':
            return parse_txt(file_path)
        
        else:
            return False, f"Unsupported file format: {file_ext}"
    
    except Exception as e:
        logger.error(f"Error parsing file {file_path}: {str(e)}")
        return False, f"Error parsing file: {str(e)}"

def parse_pdf(file_path: str) -> Tuple[bool, str]:
    """Parse PDF file and extract text content"""
    try:
        # Attempt to use PyPDF2 first
        try:
            import PyPDF2
            
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                    
            if text.strip():
                return True, text
                
            # If PyPDF2 failed to extract meaningful text, try pdfplumber
            logger.info(f"PyPDF2 extracted empty text from {file_path}, trying alternative method")
        
        except (ImportError, Exception) as e:
            logger.info(f"PyPDF2 not available or failed: {str(e)}, trying alternative method")
        
        # Try pdfplumber as a backup
        try:
            import pdfplumber
            
            with pdfplumber.open(file_path) as pdf:
                text = ""
                for page in pdf.pages:
                    text += page.extract_text() or ""
                    
            if text.strip():
                return True, text
        
        except ImportError:
            logger.warning("pdfplumber not available")
            
            # Last resort - try pdf2text if available
            try:
                from subprocess import run, PIPE
                result = run(['pdftotext', file_path, '-'], stdout=PIPE, stderr=PIPE, universal_newlines=True)
                if result.returncode == 0:
                    return True, result.stdout
                else:
                    logger.error(f"pdftotext failed: {result.stderr}")
                    return False, "Failed to extract text from PDF file"
            except Exception as e:
                logger.error(f"All PDF extraction methods failed: {str(e)}")
                return False, "Could not extract text from PDF file"
        
        # If we reach here and have no text, report failure
        if not text.strip():
            return False, "Could not extract text from PDF file"
            
        return True, text
        
    except Exception as e:
        logger.error(f"Error parsing PDF file {file_path}: {str(e)}")
        return False, f"Error parsing PDF file: {str(e)}"

def parse_docx(file_path: str) -> Tuple[bool, str]:
    """Parse DOCX file and extract text content"""
    try:
        import docx
        
        doc = docx.Document(file_path)
        fullText = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            fullText.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    fullText.append(cell.text)
        
        return True, '\n'.join(fullText)
    
    except ImportError:
        logger.warning("python-docx not available")
        return False, "python-docx library not available to parse DOCX files"
    
    except Exception as e:
        logger.error(f"Error parsing DOCX file {file_path}: {str(e)}")
        return False, f"Error parsing DOCX file: {str(e)}"

def parse_doc(file_path: str) -> Tuple[bool, str]:
    """Parse DOC file and extract text content"""
    try:
        # Try using textract if available
        try:
            import textract
            text = textract.process(file_path).decode('utf-8')
            return True, text
        except ImportError:
            logger.warning("textract not available")
        
        # Try using antiword if available
        try:
            from subprocess import run, PIPE
            result = run(['antiword', file_path], stdout=PIPE, stderr=PIPE, universal_newlines=True)
            if result.returncode == 0 and result.stdout:
                return True, result.stdout
            else:
                logger.error(f"antiword failed: {result.stderr}")
        except Exception as e:
            logger.warning(f"antiword not available or failed: {str(e)}")
        
        # If all methods fail
        return False, "Could not extract text from DOC file. Please convert to DOCX or PDF."
    
    except Exception as e:
        logger.error(f"Error parsing DOC file {file_path}: {str(e)}")
        return False, f"Error parsing DOC file: {str(e)}"

def parse_txt(file_path: str) -> Tuple[bool, str]:
    """Parse plain text file"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
            text = file.read()
        return True, text
    
    except Exception as e:
        logger.error(f"Error parsing TXT file {file_path}: {str(e)}")
        return False, f"Error parsing TXT file: {str(e)}"